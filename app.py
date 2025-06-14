import streamlit as st
import requests
import json
import hashlib
import time
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from urllib.parse import urljoin, urlparse
import re
import os
from pathlib import Path
import logging
from datetime import datetime
import uuid

# Document processing imports
import PyPDF2
import docx
from bs4 import BeautifulSoup

# Pinecone and embeddings
import pinecone
from pinecone import Pinecone, ServerlessSpec

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    # handlers=[
    #     logging.StreamHandler(),
    #     logging.FileHandler('rag_chatbot.log')
    # ]
)
logger = logging.getLogger(__name__)

@dataclass
class Document:
    """Document data structure"""
    content: str
    source: str
    title: str
    chunk_id: str
    metadata: Dict[str, Any]

class DocumentProcessor:
    """Handles processing of various document types"""
    
    def __init__(self):
        self.chunk_size = 500
        self.chunk_overlap = 50
        logger.info(f"DocumentProcessor initialized with chunk_size={self.chunk_size}, chunk_overlap={self.chunk_overlap}")
    
    def process_pdf(self, file_path: str) -> List[Document]:
        """Process PDF files"""
        documents = []
        logger.info(f"Starting PDF processing for: {file_path}")
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    full_text += page_text + "\n"
                    logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                
                logger.info(f"Total text extracted: {len(full_text)} characters")
                chunks = self._chunk_text(full_text)
                logger.info(f"Created {len(chunks)} chunks from PDF")
                
                for i, chunk in enumerate(chunks):
                    doc = Document(
                        content=chunk,
                        source=file_path,
                        title=f"PDF Document - {Path(file_path).name}",
                        chunk_id=f"pdf_{Path(file_path).stem}_{i}",
                        metadata={"type": "pdf", "page_count": len(pdf_reader.pages)}
                    )
                    documents.append(doc)
                    logger.debug(f"Created document chunk {i+1}: {len(chunk)} characters")
                
                logger.info(f"Successfully processed PDF: {len(documents)} documents created")
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}", exc_info=True)
        
        return documents
    
    def process_docx(self, file_path: str) -> List[Document]:
        """Process DOCX files"""
        documents = []
        logger.info(f"Starting DOCX processing for: {file_path}")
        try:
            doc = docx.Document(file_path)
            logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")
            full_text = ""
            for para_num, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                full_text += para_text + "\n"
                logger.debug(f"Paragraph {para_num + 1}: {len(para_text)} characters")
            
            logger.info(f"Total text extracted: {len(full_text)} characters")
            chunks = self._chunk_text(full_text)
            logger.info(f"Created {len(chunks)} chunks from DOCX")
            
            for i, chunk in enumerate(chunks):
                document = Document(
                    content=chunk,
                    source=file_path,
                    title=f"DOCX Document - {Path(file_path).name}",
                    chunk_id=f"docx_{Path(file_path).stem}_{i}",
                    metadata={"type": "docx"}
                )
                documents.append(document)
                logger.debug(f"Created document chunk {i+1}: {len(chunk)} characters")
            
            logger.info(f"Successfully processed DOCX: {len(documents)} documents created")
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}", exc_info=True)
        
        return documents
    
    def process_website(self, base_url: str, max_pages: int = 50) -> List[Document]:
        """Scrape and process website content"""
        documents = []
        visited_urls = set()
        urls_to_visit = [base_url]
        
        logger.info(f"Starting website scraping for: {base_url}, max_pages: {max_pages}")
        
        session = requests.Session()
        session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
        while urls_to_visit and len(visited_urls) < max_pages:
            url = urls_to_visit.pop(0)
            if url in visited_urls:
                continue
            
            logger.info(f"Processing URL {len(visited_urls) + 1}/{max_pages}: {url}")
            
            try:
                response = session.get(url, timeout=10)
                logger.debug(f"HTTP response: {response.status_code} for {url}")
                
                if response.status_code == 200:
                    soup = BeautifulSoup(response.content, 'html.parser')
                    
                    # Extract text content
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    text = soup.get_text()
                    text = re.sub(r'\s+', ' ', text).strip()
                    
                    logger.debug(f"Extracted {len(text)} characters from {url}")
                    
                    if len(text) > 100:  # Only process pages with substantial content
                        chunks = self._chunk_text(text)
                        logger.info(f"Created {len(chunks)} chunks from {url}")
                        
                        for i, chunk in enumerate(chunks):
                            doc = Document(
                                content=chunk,
                                source=url,
                                title=soup.title.string if soup.title else url,
                                chunk_id=f"web_{hashlib.md5(url.encode()).hexdigest()}_{i}",
                                metadata={"type": "website", "url": url}
                            )
                            documents.append(doc)
                            logger.debug(f"Created web chunk {i+1}: {len(chunk)} characters")
                    else:
                        logger.warning(f"Skipping {url}: insufficient content ({len(text)} characters)")
                    
                    # Find more URLs to scrape (only within same domain)
                    if len(visited_urls) < max_pages:
                        links = soup.find_all('a', href=True)
                        new_urls_found = 0
                        for link in links[:10]:  # Limit links per page
                            href = link['href']
                            full_url = urljoin(url, href)
                            if (urlparse(full_url).netloc == urlparse(base_url).netloc and 
                                full_url not in visited_urls and 
                                full_url not in urls_to_visit):
                                urls_to_visit.append(full_url)
                                new_urls_found += 1
                        logger.debug(f"Found {new_urls_found} new URLs from {url}")
                else:
                    logger.warning(f"HTTP {response.status_code} for {url}")
                
                visited_urls.add(url)
                time.sleep(1)  # Be respectful to the server
                
            except Exception as e:
                logger.error(f"Error processing URL {url}: {e}")
                visited_urls.add(url)
        
        logger.info(f"Website scraping completed: {len(documents)} total documents from {len(visited_urls)} pages")
        return documents
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        logger.debug(f"Chunking text with {len(words)} words")
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk = ' '.join(chunk_words)
            if len(chunk.strip()) > 50:  # Only keep substantial chunks
                chunks.append(chunk.strip())
        
        logger.debug(f"Created {len(chunks)} chunks")
        return chunks

class PineconeVectorStore:
    """Pinecone vector database manager with Pinecone embeddings"""
    
    def __init__(self, api_key: str, index_name: str = "support-docs"):
        logger.info(f"Initializing PineconeVectorStore with index: {index_name}")
        self.pc = Pinecone(api_key=api_key)
        self.index_name = index_name
        self.embedding_model = "multilingual-e5-large"
        self.dimension = 1024  # multilingual-e5-large dimension
        self.embedding_batch_size = 50  # Reduced from 100 to stay under 96 limit
        self.index = None
        logger.info(f"Using embedding model: {self.embedding_model}, dimension: {self.dimension}")
        self._setup_index()
    
    def _setup_index(self):
        """Setup Pinecone index"""
        try:
            logger.info("Setting up Pinecone index...")
            # Check if index exists
            existing_indexes = [index.name for index in self.pc.list_indexes()]
            logger.info(f"Existing indexes: {existing_indexes}")
            
            if self.index_name not in existing_indexes:
                logger.info(f"Creating new Pinecone index: {self.index_name}")
                # Create new index
                self.pc.create_index(
                    name=self.index_name,
                    dimension=self.dimension,
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
                logger.info(f"Created new Pinecone index: {self.index_name}")
                time.sleep(10)  # Wait for index to be ready
            else:
                logger.info(f"Index {self.index_name} already exists")
            
            # Connect to index
            self.index = self.pc.Index(self.index_name)
            logger.info(f"Connected to Pinecone index: {self.index_name}")
            
        except Exception as e:
            logger.error(f"Error setting up Pinecone index: {e}", exc_info=True)
            raise
    
    def _get_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings using Pinecone's embedding service with proper batching"""
        logger.info(f"Generating embeddings for {len(texts)} texts")
        try:
            all_embeddings = []
            
            # Much smaller batch size to avoid rate limits
            actual_batch_size = min(50, self.embedding_batch_size)
            logger.info(f"Using batch size: {actual_batch_size}")
            
            # Process in smaller batches to respect API limits
            for i in range(0, len(texts), actual_batch_size):
                batch_texts = texts[i:i + actual_batch_size]
                batch_num = i // actual_batch_size + 1
                total_batches = (len(texts) - 1) // actual_batch_size + 1
                
                logger.info(f"Processing embedding batch {batch_num}/{total_batches} ({len(batch_texts)} texts)")
                
                # Add retry logic for rate limiting
                max_retries = 3
                retry_delay = 5  # Start with longer delay
                
                for attempt in range(max_retries):
                    try:
                        logger.debug(f"Embedding attempt {attempt + 1} for batch {batch_num}")
                        response = self.pc.inference.embed(
                            model=self.embedding_model,
                            inputs=batch_texts,
                            parameters={"input_type": "passage", "truncate": "END"}
                        )
                        
                        # Fix: Handle EmbeddingsList response format
                        if hasattr(response, 'data') and response.data:
                            # Extract embeddings from EmbeddingsList.data
                            batch_embeddings = [embedding.values for embedding in response.data]
                            logger.debug(f"Successfully generated {len(batch_embeddings)} embeddings")
                        else:
                            # Fallback for unexpected response format
                            logger.error(f"Unexpected response format: {type(response)}")
                            batch_embeddings = [[0.0] * self.dimension for _ in batch_texts]
                        
                        all_embeddings.extend(batch_embeddings)
                        
                        # Longer delay between batches to respect rate limits
                        if i + actual_batch_size < len(texts):
                            logger.debug(f"Waiting 2 seconds before next batch...")
                            time.sleep(2)  # Increased delay
                        break
                        
                    except Exception as e:
                        error_str = str(e)
                        if "429" in error_str or "rate limit" in error_str.lower():
                            # Rate limit hit - use exponential backoff
                            if attempt < max_retries - 1:
                                logger.warning(f"Rate limit hit, retrying in {retry_delay}s (attempt {attempt + 1})")
                                time.sleep(retry_delay)
                                retry_delay *= 2  # Exponential backoff
                            else:
                                logger.error(f"Rate limit exceeded after {max_retries} attempts")
                                # Return dummy embeddings for failed batch
                                dummy_embeddings = [[0.0] * self.dimension for _ in batch_texts]
                                all_embeddings.extend(dummy_embeddings)
                        else:
                            if attempt < max_retries - 1:
                                logger.warning(f"Embedding attempt {attempt + 1} failed, retrying in {retry_delay}s: {e}")
                                time.sleep(retry_delay)
                                retry_delay *= 2
                            else:
                                logger.error(f"Failed to generate embeddings after {max_retries} attempts: {e}")
                                # Return dummy embeddings for failed batch
                                dummy_embeddings = [[0.0] * self.dimension for _ in batch_texts]
                                all_embeddings.extend(dummy_embeddings)
            
            logger.info(f"Successfully generated {len(all_embeddings)} embeddings")
            return all_embeddings
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}", exc_info=True)
            # Fallback to dummy embeddings for testing
            return [[0.0] * self.dimension for _ in texts]
    
    def add_documents(self, documents: List[Document]):
        """Add documents to Pinecone vector store"""
        if not self.index:
            logger.error("Pinecone index not initialized")
            return
        
        logger.info(f"Adding {len(documents)} documents to Pinecone")
        
        # Prepare texts for embedding
        texts = [doc.content for doc in documents]
        
        # Generate embeddings with proper batching
        st.info(f"Generating embeddings for {len(texts)} documents in batches of {self.embedding_batch_size}...")
        embeddings = self._get_embeddings(texts)
        
        # Prepare vectors for upsert
        vectors_to_upsert = []
        for doc, embedding in zip(documents, embeddings):
            vector_id = str(uuid.uuid4())  # Generate unique ID
            vector = {
                "id": vector_id,
                "values": embedding,
                "metadata": {
                    "content": doc.content,
                    "source": doc.source,
                    "title": doc.title,
                    "chunk_id": doc.chunk_id,
                    "doc_type": doc.metadata.get("type", "unknown"),
                    "url": doc.metadata.get("url", ""),
                }
            }
            vectors_to_upsert.append(vector)
        
        logger.info(f"Prepared {len(vectors_to_upsert)} vectors for upsert")
        
        # Upsert vectors in batches
        upsert_batch_size = 100
        total_batches = (len(vectors_to_upsert) - 1) // upsert_batch_size + 1
        
        for i in range(0, len(vectors_to_upsert), upsert_batch_size):
            batch = vectors_to_upsert[i:i + upsert_batch_size]
            batch_num = i // upsert_batch_size + 1
            
            try:
                logger.info(f"Upserting batch {batch_num}/{total_batches} ({len(batch)} vectors)")
                self.index.upsert(vectors=batch)
                logger.info(f"Successfully upserted batch {batch_num}")
                st.info(f"Uploaded batch {batch_num} of {total_batches}")
            except Exception as e:
                logger.error(f"Error upserting batch {batch_num}: {e}", exc_info=True)
        
        logger.info(f"Successfully added {len(documents)} documents to Pinecone")
    
    def search(self, query: str, limit: int = 5) -> List[Dict]:
        """Search for relevant documents"""
        if not self.index:
            logger.error("Pinecone index not initialized")
            return []
        
        logger.info(f"Searching Pinecone for query: '{query[:50]}...' (limit: {limit})")
        
        try:
            # Generate query embedding
            logger.debug("Generating query embedding...")
            query_embedding = self._get_embeddings([query])[0]
            
            # Search in Pinecone
            logger.debug("Performing vector search...")
            search_result = self.index.query(
                vector=query_embedding,
                top_k=limit,
                include_metadata=True,
                filter=None  # Add filters if needed
            )
            
            results = []
            for match in search_result['matches']:
                score = match['score']
                logger.debug(f"Match found with score: {score}")
                
                if score >= 0.3:  # Minimum similarity threshold
                    results.append({
                        "content": match['metadata']['content'],
                        "source": match['metadata']['source'],
                        "title": match['metadata']['title'],
                        "score": score,
                        "metadata": {
                            "type": match['metadata']['doc_type'],
                            "url": match['metadata'].get('url', ''),
                            "chunk_id": match['metadata']['chunk_id']
                        }
                    })
                else:
                    logger.debug(f"Match below threshold ({score} < 0.3), skipping")
            
            logger.info(f"Found {len(results)} relevant matches above threshold")
            return results
            
        except Exception as e:
            logger.error(f"Error searching Pinecone: {e}", exc_info=True)
            return []
    
    def get_index_stats(self) -> Dict:
        """Get index statistics"""
        if not self.index:
            logger.warning("Index not initialized for stats retrieval")
            return {"error": "Index not initialized"}
        
        try:
            logger.debug("Fetching index statistics...")
            stats = self.index.describe_index_stats()
            result = {
                "total_vectors": stats.get('total_vector_count', 0),
                "dimension": stats.get('dimension', 0),
                "index_fullness": stats.get('index_fullness', 0)
            }
            logger.info(f"Index stats: {result}")
            return result
        except Exception as e:
            logger.error(f"Error getting index stats: {e}", exc_info=True)
            return {"error": str(e)}

class RAGChatbot:
    """Main RAG chatbot class with Pinecone integration"""
    
    def __init__(self, openrouter_api_key: str, pinecone_api_key: str, model: str = "anthropic/claude-3-haiku"):
        logger.info(f"Initializing RAGChatbot with model: {model}")
        self.openrouter_api_key = openrouter_api_key
        self.pinecone_api_key = pinecone_api_key
        self.model = model
        self.vector_store = PineconeVectorStore(pinecone_api_key)
        self.doc_processor = DocumentProcessor()
        self.base_url = "https://openrouter.ai/api/v1/chat/completions"
        logger.info("RAGChatbot initialization completed")
    
    def load_documents(self, sources: Dict[str, Any]):
        """Load documents from various sources"""
        logger.info(f"Loading documents from sources: {list(sources.keys())}")
        all_documents = []
        
        # Process website
        if "website" in sources and sources["website"]:
            st.info("🌐 Scraping website content...")
            logger.info(f"Processing website: {sources['website']}")
            website_docs = self.doc_processor.process_website(
                sources["website"], 
                max_pages=sources.get("max_pages", 20)
            )
            all_documents.extend(website_docs)
            st.success(f"✅ Processed {len(website_docs)} website chunks")
            logger.info(f"Website processing completed: {len(website_docs)} chunks")
        
        # Process uploaded files
        if "files" in sources:
            logger.info(f"Processing {len(sources['files'])} uploaded files")
            for uploaded_file in sources["files"]:
                file_path = f"temp_{uploaded_file.name}"
                logger.info(f"Processing uploaded file: {uploaded_file.name}")
                
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                if uploaded_file.name.endswith('.pdf'):
                    st.info(f"📄 Processing PDF: {uploaded_file.name}")
                    pdf_docs = self.doc_processor.process_pdf(file_path)
                    all_documents.extend(pdf_docs)
                    st.success(f"✅ Processed {len(pdf_docs)} PDF chunks")
                    logger.info(f"PDF processing completed: {len(pdf_docs)} chunks")
                
                elif uploaded_file.name.endswith('.docx'):
                    st.info(f"📝 Processing DOCX: {uploaded_file.name}")
                    docx_docs = self.doc_processor.process_docx(file_path)
                    all_documents.extend(docx_docs)
                    st.success(f"✅ Processed {len(docx_docs)} DOCX chunks")
                    logger.info(f"DOCX processing completed: {len(docx_docs)} chunks")
                
                # Clean up temp file
                os.remove(file_path)
                logger.debug(f"Cleaned up temporary file: {file_path}")
        
        # Add to Pinecone vector store
        if all_documents:
            st.info("🔍 Creating embeddings with Pinecone and storing in vector database...")
            logger.info(f"Adding {len(all_documents)} documents to vector store")
            self.vector_store.add_documents(all_documents)
            st.success(f"✅ Successfully loaded {len(all_documents)} document chunks into Pinecone!")
            logger.info(f"Document loading completed: {len(all_documents)} total chunks")
        else:
            logger.warning("No documents to process")
        
        return len(all_documents)
    
    def generate_response(self, query: str, relevant_docs: List[Dict]) -> str:
        """Generate response using OpenRouter API"""
        logger.info(f"Generating response for query: '{query[:50]}...' with {len(relevant_docs)} relevant docs")
        
        if not relevant_docs:
            logger.info("No relevant documents found, returning default response")
            return "I don't know. I couldn't find relevant information in the support documentation to answer your question."
        
        # Prepare context from relevant documents
        context = "\n\n".join([
            f"Source: {doc['title']}\nContent: {doc['content']}"
            for doc in relevant_docs
        ])
        
        logger.debug(f"Context length: {len(context)} characters")
        
        system_prompt = """You are a helpful customer support assistant. You can ONLY answer questions based on the provided support documentation context. 

IMPORTANT RULES:
1. ONLY use information from the provided context to answer questions
2. If the answer is not in the context, respond with "I don't know"
3. Be helpful and provide detailed answers when information is available
4. Always be polite and professional
5. If you reference information, mention the source when helpful

Context from support documentation:
{context}
"""
        
        messages = [
            {"role": "system", "content": system_prompt.format(context=context)},
            {"role": "user", "content": query}
        ]
        
        try:
            logger.debug(f"Sending request to OpenRouter API with model: {self.model}")
            headers = {
                "Authorization": f"Bearer {self.openrouter_api_key}",
                "Content-Type": "application/json"
            }
            
            data = {
                "model": self.model,
                "messages": messages,
                "max_tokens": 500,
                "temperature": 0.1
            }
            
            response = requests.post(self.base_url, headers=headers, json=data)
            logger.debug(f"OpenRouter API response status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                generated_response = result["choices"][0]["message"]["content"]
                logger.info(f"Successfully generated response: {len(generated_response)} characters")
                return generated_response
            else:
                logger.error(f"API Error: {response.status_code} - {response.text}")
                return "I apologize, but I'm experiencing technical difficulties. Please try again later."
                
        except Exception as e:
            logger.error(f"Error generating response: {e}", exc_info=True)
            return "I apologize, but I'm experiencing technical difficulties. Please try again later."
    
    def chat(self, query: str) -> Dict[str, Any]:
        """Main chat function"""
        logger.info(f"Processing chat query: '{query[:50]}...'")
        
        # Search for relevant documents in Pinecone
        relevant_docs = self.vector_store.search(query, limit=5)
        
        # Generate response
        response = self.generate_response(query, relevant_docs)
        
        result = {
            "response": response,
            "sources": relevant_docs,
            "timestamp": datetime.now().isoformat()
        }
        
        logger.info(f"Chat processing completed: {len(response)} character response with {len(relevant_docs)} sources")
        return result
    
    def get_stats(self) -> Dict:
        """Get system statistics"""
        logger.debug("Fetching system statistics")
        return self.vector_store.get_index_stats()

def main():
    """Streamlit app main function"""
    st.set_page_config(
        page_title="🤖 Pinecone RAG Customer Support Chatbot",
        page_icon="🤖",
        layout="wide"
    )
    
    st.title("🤖 Advanced RAG Customer Support Chatbot")
    st.markdown("*Powered by Pinecone Vector Database & Embeddings*")
    st.markdown("---")
    
    # Get default API keys from environment
    default_openrouter_key = os.getenv("OPENROUTER_API_KEY")
    default_pinecone_key = os.getenv("PINECONE_API_KEY")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Keys input
        st.subheader("🔑 API Keys")
        
        # Use default keys toggle
        use_default_keys = st.checkbox(
            "🔧 Use Default Keys", 
            value=bool(default_openrouter_key and default_pinecone_key),
            help="Use API keys from environment variables (.env file)"
        )
        
        if use_default_keys and default_openrouter_key and default_pinecone_key:
            st.success("✅ Using default API keys from environment")
            openrouter_api_key = default_openrouter_key
            pinecone_api_key = default_pinecone_key
            
            # Show masked keys for verification
            st.text(f"OpenRouter: {default_openrouter_key[:8]}...{default_openrouter_key[-4:]}")
            st.text(f"Pinecone: {default_pinecone_key[:8]}...{default_pinecone_key[-4:]}")
        else:
            if use_default_keys:
                st.warning("⚠️ Default keys not found in environment. Please provide keys manually.")
            
            openrouter_api_key = st.text_input(
                "OpenRouter API Key",
                type="password",
                value=default_openrouter_key if default_openrouter_key else "",
                help="Get your API key from https://openrouter.ai/"
            )
            
            pinecone_api_key = st.text_input(
                "Pinecone API Key",
                type="password",
                value=default_pinecone_key if default_pinecone_key else "",
                help="Get your API key from https://pinecone.io/"
            )
        
        # Model selection
        model = st.selectbox(
            "Select Model",
            [
                "meta-llama/llama-3.3-8b-instruct:free",
                "qwen/qwen3-8b:free",
                "google/gemma-3-4b-it:free",
                "mistralai/mistral-small-3.1-24b-instruct:free"
            ]
        )
        
        st.markdown("---")
        st.header("📚 Document Sources")
        
        # Check existing index status if we have API key
        existing_docs_count = 0
        if pinecone_api_key:
            try:
                temp_vector_store = PineconeVectorStore(pinecone_api_key)
                stats = temp_vector_store.get_index_stats()
                existing_docs_count = stats.get('total_vectors', 0)
                
                if existing_docs_count > 0:
                    st.info(f"📊 Found {existing_docs_count} existing documents in vector database")
                    
                    # Option to use existing index or re-index
                    index_option = st.radio(
                        "Index Options",
                        ["Use Existing Index", "Re-index Documents"],
                        help="Choose whether to use existing documents or load new ones"
                    )
                    
                    if index_option == "Use Existing Index":
                        st.success("✅ Using existing indexed documents")
                        # Set flag to skip document loading
                        st.session_state.use_existing_index = True
                        st.session_state.docs_loaded = True
                        st.session_state.doc_count = existing_docs_count
                    else:
                        st.warning("⚠️ Will re-index documents (this will replace existing ones)")
                        st.session_state.use_existing_index = False
                else:
                    st.info("📝 No existing documents found. Ready to index new documents.")
                    st.session_state.use_existing_index = False
                    
            except Exception as e:
                logger.warning(f"Could not check existing index: {e}")
                st.warning("⚠️ Could not check existing index status")
                st.session_state.use_existing_index = False
        
        # Only show document loading options if not using existing index
        if not st.session_state.get('use_existing_index', False):
            # Website URL
            website_url = st.text_input(
                "Website URL",
                value="https://www.angelone.in/support",
                help="Base URL to scrape for support documentation"
            )
            
            max_pages = st.slider("Max pages to scrape", 5, 100, 20)
            
            # File uploads
            uploaded_files = st.file_uploader(
                "Upload Documents",
                type=['pdf', 'docx'],
                accept_multiple_files=True,
                help="Upload PDF or DOCX files containing support documentation"
            )
            
            # Load documents button
            load_docs = st.button("🔄 Load Documents", type="primary")
        else:
            st.markdown("*Using existing indexed documents*")
            load_docs = False
    
    # Check API keys
    if not openrouter_api_key or not pinecone_api_key:
        st.warning("⚠️ Please enter both OpenRouter and Pinecone API keys in the sidebar to continue.")
        
        # Show different messages based on whether default keys are available
        if default_openrouter_key and default_pinecone_key:
            st.info("💡 **Quick Start:** Enable 'Use Default Keys' in the sidebar to get started immediately!")
        else:
            st.info("📖 **Getting Started:**")
            st.markdown("""
            1. **OpenRouter API Key**: Sign up at [https://openrouter.ai/](https://openrouter.ai/)
            2. **Pinecone API Key**: Sign up at [https://pinecone.io/](https://pinecone.io/)
            3. Enter both keys in the sidebar (or add them to .env file)
            4. Load your support documents
            5. Start chatting!
            """)
        st.stop()
    
    # Initialize chatbot
    if 'chatbot' not in st.session_state:
        try:
            with st.spinner("Initializing Pinecone connection..."):
                st.session_state.chatbot = RAGChatbot(openrouter_api_key, pinecone_api_key, model)
            st.success("✅ Connected to Pinecone!")
            
            # If using existing index, show the count
            if st.session_state.get('use_existing_index', False):
                st.success(f"✅ Using existing index with {existing_docs_count} documents!")
                
        except Exception as e:
            st.error(f"❌ Error initializing Pinecone: {e}")
            st.stop()
    
    # Update chatbot settings if changed
    if 'chatbot' in st.session_state:
        st.session_state.chatbot.model = model
        st.session_state.chatbot.openrouter_api_key = openrouter_api_key
    
    # Load documents (only if not using existing index)
    if load_docs and not st.session_state.get('use_existing_index', False):
        sources = {
            "website": website_url if website_url else None,
            "max_pages": max_pages,
            "files": uploaded_files if uploaded_files else []
        }
        
        if sources["website"] or sources["files"]:
            # Show warning if re-indexing
            if existing_docs_count > 0:
                st.warning(f"⚠️ Re-indexing will replace {existing_docs_count} existing documents!")
                if not st.button("⚠️ Confirm Re-indexing", type="secondary"):
                    st.stop()
            
            with st.spinner("Loading and processing documents..."):
                try:
                    doc_count = st.session_state.chatbot.load_documents(sources)
                    st.session_state.docs_loaded = True
                    st.session_state.doc_count = doc_count
                    st.session_state.use_existing_index = False
                except Exception as e:
                    st.error(f"Error loading documents: {e}")
        else:
            st.warning("⚠️ Please provide at least a website URL or upload some files.")
    
    # Main chat interface
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("💬 Chat Interface")
        
        # Check if documents are loaded or using existing index
        if not st.session_state.get('docs_loaded', False):
            st.info("👆 Please load your support documents using the sidebar first.")
        else:
            if st.session_state.get('use_existing_index', False):
                st.success(f"✅ Using existing index with {st.session_state.doc_count} document chunks!")
            else:
                st.success(f"✅ {st.session_state.doc_count} document chunks loaded and ready!")
            
            # Initialize chat history
            if 'messages' not in st.session_state:
                st.session_state.messages = []
            
            # Display chat history in a scrollable container
            chat_container = st.container()
            with chat_container:
                if st.session_state.messages:
                    for message in st.session_state.messages:
                        with st.chat_message(message["role"]):
                            st.markdown(message["content"])
                            
                            # Show sources for assistant messages
                            if message["role"] == "assistant" and "sources" in message:
                                sources = message["sources"]
                                if sources and not message["content"].startswith("I don't know"):
                                    with st.expander("📚 Sources", expanded=False):
                                        for i, source in enumerate(sources[:3]):
                                            st.markdown(f"**Source {i+1}:** {source['title']}")
                                            st.markdown(f"*Similarity: {source['score']:.3f}*")
                                            st.markdown(f"*Type: {source['metadata']['type']}*")
                                            st.markdown(f"```\n{source['content'][:200]}...\n```")
                                            st.markdown("---")
                else:
                    st.markdown("*Start a conversation by asking a question below...*")
            
            # Chat input - always at the bottom
            prompt = st.chat_input("Ask me anything about the support documentation...")
            
            # Handle new message
            if prompt:
                # Add user message to chat history
                st.session_state.messages.append({"role": "user", "content": prompt})
                
                # Generate assistant response
                with st.spinner("Searching Pinecone and generating response..."):
                    chat_result = st.session_state.chatbot.chat(prompt)
                    response = chat_result["response"]
                    sources = chat_result["sources"]
                
                # Add assistant response to chat history with sources
                st.session_state.messages.append({
                    "role": "assistant", 
                    "content": response,
                    "sources": sources
                })
                
                # Rerun to refresh the display
                st.rerun()
    
    with col2:
        st.header("📊 System Info")
        
        # API Keys status
        st.subheader("🔑 API Status")
        if use_default_keys:
            st.success("✅ Using Default Keys")
        else:
            st.info("🔧 Using Custom Keys")
        
        # Pinecone stats
        if 'chatbot' in st.session_state:
            try:
                stats = st.session_state.chatbot.get_stats()
                if 'error' not in stats:
                    st.metric("Vectors in Pinecone", stats.get('total_vectors', 0))
                    st.metric("Vector Dimension", stats.get('dimension', 0))
                    st.metric("Index Fullness", f"{stats.get('index_fullness', 0)*100:.1f}%")
                    
                    # Show index status
                    if st.session_state.get('use_existing_index', False):
                        st.success("📊 Using Existing Index")
                    else:
                        st.info("🆕 Newly Indexed")
                        
                else:
                    st.error(f"Stats error: {stats['error']}")
            except Exception as e:
                st.error(f"Error fetching stats: {e}")
        
        if st.session_state.get('docs_loaded', False):
            st.metric("Documents Loaded", st.session_state.doc_count)
            st.metric("Chat Messages", len(st.session_state.messages) if 'messages' in st.session_state else 0)
        
        # Clear index button
        if st.session_state.get('docs_loaded', False):
            st.markdown("---")
            st.subheader("🗑️ Index Management")
            if st.button("🗑️ Clear Index", type="secondary", help="Remove all documents from vector database"):
                try:
                    if 'chatbot' in st.session_state:
                        # Clear the index by deleting all vectors
                        st.session_state.chatbot.vector_store.index.delete(delete_all=True)
                        st.success("✅ Index cleared successfully!")
                        
                        # Reset session state
                        st.session_state.docs_loaded = False
                        st.session_state.use_existing_index = False
                        if 'doc_count' in st.session_state:
                            del st.session_state.doc_count
                        
                        st.rerun()
                except Exception as e:
                    st.error(f"Error clearing index: {e}")
        
        st.markdown("---")
        st.header("ℹ️ How it works")
        st.markdown("""
        1. **Document Processing**: Website scraping and file processing
        2. **Chunking**: Text split with overlap
        3. **Pinecone Embeddings**: Generate high-quality vectors
        4. **Vector Storage**: Store in Pinecone serverless index
        5. **Semantic Search**: Find relevant content
        6. **LLM Generation**: Generate contextual responses
        """)

if __name__ == "__main__":
    main()